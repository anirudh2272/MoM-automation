"""
MoM Live AI — Multilingual Minutes of Meeting Automation System

A portfolio-ready Streamlit application for end-to-end meeting intelligence:
audio upload → transcription → meeting intelligence extraction → 52-language
translation → professional DOCX export.

Tech stack:
- Streamlit (UI)
- faster-whisper (CPU, int8) for transcription
- FFmpeg for audio conversion
- deep-translator (Google Translate) for multilingual translation
- python-docx for report generation
"""

import gc
import html
import os
import re
import subprocess
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from faster_whisper import WhisperModel

try:
    from deep_translator import GoogleTranslator
    TRANSLATE_AVAILABLE = True
except Exception:
    GoogleTranslator = None
    TRANSLATE_AVAILABLE = False

try:
    from google import genai as google_genai
    GEMINI_SDK_AVAILABLE = True
except Exception:
    google_genai = None
    GEMINI_SDK_AVAILABLE = False

import json as _json


# ══════════════════════════════════════════════════════════════════
#  CONFIG
# ══════════════════════════════════════════════════════════════════
APP_NAME = "MoM Live AI"
MAX_FILE_MB = 25
MAX_TRANSLATE_CHARS = 4500
TMP_DIR = "/tmp"
SUPPORTED_EXTS = ["mp3", "wav", "m4a", "mp4"]

LANG_NAMES: Dict[str, str] = {
    "af": "Afrikaans",    "ar": "Arabic",       "az": "Azerbaijani",  "bn": "Bengali",
    "cs": "Czech",        "de": "German",       "en": "English",      "es": "Spanish",
    "et": "Estonian",     "fa": "Persian",      "fi": "Finnish",      "fr": "French",
    "gl": "Galician",     "gu": "Gujarati",     "he": "Hebrew",       "hi": "Hindi",
    "hr": "Croatian",     "id": "Indonesian",   "it": "Italian",      "ja": "Japanese",
    "ka": "Georgian",     "kk": "Kazakh",       "km": "Khmer",        "ko": "Korean",
    "lt": "Lithuanian",   "lv": "Latvian",      "mk": "Macedonian",   "ml": "Malayalam",
    "mn": "Mongolian",    "mr": "Marathi",      "my": "Burmese",      "ne": "Nepali",
    "nl": "Dutch",        "pl": "Polish",       "ps": "Pashto",       "pt": "Portuguese",
    "ro": "Romanian",     "ru": "Russian",      "si": "Sinhala",      "sl": "Slovenian",
    "sv": "Swedish",      "sw": "Swahili",      "ta": "Tamil",        "te": "Telugu",
    "th": "Thai",         "tl": "Filipino",     "tr": "Turkish",      "uk": "Ukrainian",
    "ur": "Urdu",         "vi": "Vietnamese",   "xh": "Xhosa",        "zh": "Chinese",
}

# Whisper "base" struggles with these — recommend "Better" mode
NEEDS_BETTER_MODE = {
    "te", "hi", "ta", "ml", "mr", "gu", "bn", "ur", "si", "ne",
    "ar", "fa", "ps", "he", "ka", "my", "km", "th", "mn",
    "sw", "af", "xh", "tl", "az", "kk", "vi",
}

# Code remapping for Google Translate
GT_CODE_MAP = {"zh": "zh-CN", "he": "iw"}

# Gemini model for AI-powered meeting analysis (free-tier compatible)
GEMINI_MODEL = "gemini-2.5-flash"


# ══════════════════════════════════════════════════════════════════
#  PAGE CONFIG & STYLING
# ══════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title=f"{APP_NAME} — Multilingual Meeting Intelligence",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

/* Apple uses SF Pro on Apple devices; Inter is the closest cross-platform fallback */
:root {
    --sf: -apple-system, BlinkMacSystemFont, "SF Pro Display", "SF Pro Text",
          "Inter", "Helvetica Neue", Helvetica, Arial, sans-serif;
    --ink: #1d1d1f;          /* Apple near-black text */
    --ink-soft: #6e6e73;     /* Apple secondary gray */
    --line: #d2d2d7;         /* Apple hairline */
    --bg: #ffffff;           /* pure white */
    --bg-soft: #f5f5f7;      /* Apple light gray panel */
    --blue: #0071e3;         /* Apple blue */
    --blue-hover: #0077ed;
    --radius: 18px;
}

*, *::before, *::after { box-sizing: border-box; }
html, body, [class*="css"] {
    font-family: var(--sf);
    -webkit-font-smoothing: antialiased;
    -moz-osx-font-smoothing: grayscale;
    letter-spacing: -0.01em;
}

/* ── Base canvas: clean white, no gradients ── */
.stApp { background: var(--bg); }
.main .block-container {
    max-width: 980px;
    padding: 1.5rem 1.5rem 5rem;
}

/* ── Typography ── */
p, li { color: var(--ink-soft); line-height: 1.5; font-weight: 400; }
span, label { color: var(--ink); }
h1, h2, h3, h4 { color: var(--ink); font-weight: 600; letter-spacing: -0.02em; }
strong, b { color: var(--ink); font-weight: 600; }
hr { border: none; height: 1px; background: var(--line); margin: 3rem 0; }

/* ── Hero ── */
.hero { text-align: center; padding: 3rem 0 2.5rem; }
.eyebrow {
    display: inline-block;
    font-size: .8rem;
    font-weight: 600;
    letter-spacing: -0.01em;
    color: var(--blue);
    margin-bottom: .75rem;
    text-transform: none;
}
.hero-title {
    font-family: var(--sf);
    font-size: clamp(2.6rem, 6vw, 4rem);
    font-weight: 700;
    letter-spacing: -0.03em;
    line-height: 1.05;
    color: var(--ink);
    margin: 0 0 1rem;
}
.hero-title em {
    font-style: normal;
    color: var(--blue);
}
.hero-sub {
    color: var(--ink-soft) !important;
    font-size: 1.25rem;
    font-weight: 400;
    line-height: 1.4;
    max-width: 620px;
    margin: 0 auto;
    letter-spacing: -0.01em;
}

/* ── Stat cards: flat gray, large radius ── */
.stat-card {
    background: var(--bg-soft);
    border: none;
    border-radius: var(--radius);
    padding: 1.6rem 1rem;
    text-align: center;
    transition: transform .25s ease;
}
.stat-card:hover { transform: scale(1.015); }
.stat-icon { font-size: 1.5rem; display: block; margin-bottom: .5rem; }
.stat-value {
    font-family: var(--sf);
    font-size: 2rem;
    font-weight: 600;
    color: var(--ink);
    line-height: 1;
    letter-spacing: -0.02em;
}
.stat-label {
    font-size: .78rem;
    font-weight: 500;
    color: var(--ink-soft);
    margin-top: .5rem;
    letter-spacing: -0.01em;
}

/* ── Section headers: simple, bold, no numbered markers ── */
.sec-hdr {
    font-family: var(--sf);
    font-size: 1.9rem;
    font-weight: 600;
    color: var(--ink);
    margin: 3rem 0 1.25rem;
    letter-spacing: -0.02em;
    display: flex;
    align-items: baseline;
    gap: .6rem;
}
.sec-num {
    font-size: .9rem;
    font-weight: 600;
    color: var(--ink-soft);
}

/* ── Cards ── */
.glass-card {
    background: var(--bg-soft);
    border: none;
    border-radius: var(--radius);
    padding: 1.5rem 1.6rem;
    color: var(--ink);
    line-height: 1.55;
    font-size: 1.02rem;
    white-space: pre-wrap;
    word-break: break-word;
    letter-spacing: -0.01em;
}

/* ── Result boxes: flat gray ── */
.result-box {
    background: var(--bg-soft);
    border: none;
    border-radius: var(--radius);
    padding: 1.5rem 1rem;
    text-align: center;
    transition: transform .25s ease;
}
.result-box:hover { transform: scale(1.015); }
.result-val {
    font-family: var(--sf);
    font-size: 2.4rem;
    font-weight: 600;
    color: var(--ink);
    line-height: 1;
    letter-spacing: -0.025em;
}
.result-lbl {
    font-size: .78rem;
    font-weight: 500;
    color: var(--ink-soft);
    margin-top: .6rem;
    letter-spacing: -0.01em;
}

/* ── Pills: subtle, rounded ── */
.pill {
    display: inline-flex;
    align-items: center;
    gap: .35rem;
    border-radius: 980px;
    padding: .4rem .9rem;
    font-size: .82rem;
    font-weight: 500;
    margin: .2rem .25rem .2rem 0;
    background: var(--bg-soft);
    color: var(--ink);
    border: none;
    letter-spacing: -0.01em;
}
.pill-v, .pill-p, .pill-g, .pill-a {
    background: var(--bg-soft);
    color: var(--ink);
    border: none;
}

/* ── Empty state ── */
.empty-state {
    padding: 5rem 2rem;
    text-align: center;
    border: none;
    border-radius: 28px;
    background: var(--bg-soft);
    margin-top: 1rem;
}
.empty-icon { font-size: 4rem; display: block; margin-bottom: 1.25rem; }
.empty-title {
    font-family: var(--sf);
    font-size: 1.9rem;
    font-weight: 600;
    color: var(--ink);
    margin-bottom: .5rem;
    letter-spacing: -0.02em;
}
.empty-text { color: var(--ink-soft) !important; font-size: 1.05rem; }

/* ── Sidebar: light gray ── */
section[data-testid="stSidebar"] {
    background: var(--bg-soft) !important;
    border-right: 1px solid var(--line) !important;
}
section[data-testid="stSidebar"] * { color: var(--ink) !important; }
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
    color: var(--ink) !important;
    font-family: var(--sf) !important;
    font-weight: 600 !important;
    font-style: normal !important;
    letter-spacing: -0.02em !important;
}
section[data-testid="stSidebar"] strong { color: var(--ink) !important; font-weight: 600 !important; }
section[data-testid="stSidebar"] .stCaption,
section[data-testid="stSidebar"] small { color: var(--ink-soft) !important; }

/* ── Buttons: solid Apple blue, pill, no gradient/glow ── */
.stButton > button {
    background: var(--blue) !important;
    color: #fff !important;
    border: none !important;
    border-radius: 980px !important;
    font-family: var(--sf) !important;
    font-weight: 500 !important;
    font-size: 1.05rem !important;
    width: 100% !important;
    padding: .85rem 1.8rem !important;
    box-shadow: none !important;
    letter-spacing: -0.01em !important;
    transition: background .2s ease !important;
}
.stButton > button:hover {
    background: var(--blue-hover) !important;
    transform: none !important;
}
.stDownloadButton > button {
    background: var(--blue) !important;
    color: #fff !important;
    border: none !important;
    border-radius: 980px !important;
    font-weight: 500 !important;
    font-size: 1.05rem !important;
    width: 100% !important;
    padding: .85rem 1.8rem !important;
    box-shadow: none !important;
    letter-spacing: -0.01em !important;
    transition: background .2s ease !important;
}
.stDownloadButton > button:hover { background: var(--blue-hover) !important; transform: none !important; }

/* ── File uploader: clean gray ── */
.stFileUploader > div {
    background: var(--bg-soft) !important;
    border: 1px dashed var(--line) !important;
    border-radius: var(--radius) !important;
    transition: border-color .2s !important;
}
.stFileUploader > div:hover { border-color: var(--blue) !important; }

/* ── Inputs ── */
.stSelectbox > div > div, .stTextInput > div > div {
    background: #fff !important;
    border: 1px solid var(--line) !important;
    border-radius: 12px !important;
    color: var(--ink) !important;
}
.stTextInput > div > div:focus-within,
.stSelectbox > div > div:focus-within {
    border-color: var(--blue) !important;
    box-shadow: 0 0 0 3px rgba(0,113,227,.15) !important;
}

/* ── Tabs: pill segmented control ── */
.stTabs [data-baseweb="tab-list"] {
    background: var(--bg-soft);
    border-radius: 980px;
    padding: 4px;
    border: none;
    gap: 4px;
}
.stTabs [data-baseweb="tab"] {
    color: var(--ink) !important;
    border-radius: 980px !important;
    padding: .45rem 1.4rem !important;
    font-weight: 500 !important;
    letter-spacing: -0.01em !important;
}
.stTabs [aria-selected="true"] {
    background: #fff !important;
    color: var(--ink) !important;
    box-shadow: 0 1px 3px rgba(0,0,0,.08) !important;
}

/* ── Expander ── */
.streamlit-expanderHeader {
    background: var(--bg-soft) !important;
    border: none !important;
    border-radius: 12px !important;
    color: var(--ink) !important;
    font-weight: 500 !important;
}
.streamlit-expanderHeader:hover { background: #ececf0 !important; }
.streamlit-expanderContent {
    background: #fff !important;
    border: 1px solid var(--line) !important;
    border-top: none !important;
    border-radius: 0 0 12px 12px !important;
}

/* ── Alerts ── */
.stAlert {
    background: var(--bg-soft) !important;
    border: 1px solid var(--line) !important;
    border-radius: 14px !important;
    color: var(--ink) !important;
}

/* ── Progress: Apple blue ── */
.stProgress > div > div { background: var(--blue) !important; border-radius: 980px !important; }
.stProgress > div { background: #e8e8ed !important; border-radius: 980px !important; }

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 8px; height: 8px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: #c7c7cc; border-radius: 980px; }
::-webkit-scrollbar-thumb:hover { background: #aeaeb2; }

/* ── Audio player ── */
audio { width: 100%; }

#MainMenu, footer, header { visibility: hidden; }
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
#  UTILITIES
# ══════════════════════════════════════════════════════════════════
def safe_filename(name: str) -> str:
    """Make a filename safe for filesystem use."""
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", name)[:120] or "file"


def clean_docx(text: str) -> str:
    """Strip illegal XML control characters that break DOCX writers."""
    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", str(text or ""))


def esc(text: str) -> str:
    """HTML-escape user-facing strings before inserting into custom HTML cards."""
    return html.escape(text or "")


def save_uploaded_file(uploaded_file) -> str:
    """Persist the uploaded file to /tmp and return its path."""
    safe = safe_filename(uploaded_file.name)
    path = os.path.join(TMP_DIR, f"{int(time.time())}_{safe}")
    with open(path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return path


# ══════════════════════════════════════════════════════════════════
#  AUDIO CONVERSION
# ══════════════════════════════════════════════════════════════════
def convert_to_wav(input_path: str) -> Tuple[str, Optional[str]]:
    """
    Convert any uploaded audio/video to 16kHz mono PCM WAV using FFmpeg.
    On failure, fall back to the original file (faster-whisper can often read it).
    Returns (path_to_use, warning_message_or_None).
    """
    output_path = str(Path(input_path).with_suffix("")) + "_16k.wav"
    cmd = [
        "ffmpeg", "-y", "-i", input_path,
        "-vn", "-acodec", "pcm_s16le",
        "-ar", "16000", "-ac", "1",
        "-loglevel", "error",
        output_path,
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
        if result.returncode != 0:
            return input_path, f"FFmpeg conversion failed; using original file."
        if not os.path.exists(output_path) or os.path.getsize(output_path) < 1000:
            return input_path, "Converted file was empty; using original file."
        return output_path, None
    except FileNotFoundError:
        return input_path, "FFmpeg not installed. Add 'ffmpeg' to packages.txt."
    except subprocess.TimeoutExpired:
        return input_path, "FFmpeg timed out; using original file."
    except Exception as exc:
        return input_path, f"FFmpeg error: {str(exc)[:150]}"


# ══════════════════════════════════════════════════════════════════
#  WHISPER (cached so model loads once per session)
# ══════════════════════════════════════════════════════════════════
@st.cache_resource(show_spinner=False)
def load_whisper_model(size: str) -> WhisperModel:
    """Load a faster-whisper model on CPU with int8 quantization."""
    return WhisperModel(size, device="cpu", compute_type="int8")


def transcribe_audio(
    audio_path: str,
    model_size: str,
    language_hint: str = "auto",
) -> Tuple[str, List[Dict], str, float, float]:
    """
    Transcribe audio with two passes:
      Pass 1 — Standard (beam=5, vad off, no_speech_threshold=0.1)
      Pass 2 — Ultra-lenient fallback if nothing was returned
    Returns: (transcript, segments, detected_language, lang_probability_pct, duration_sec)
    """
    model = load_whisper_model(model_size)
    language = None if language_hint == "auto" else language_hint

    def _run(beam: int, no_speech: float, log_prob: Optional[float] = None):
        kwargs = dict(
            beam_size=beam,
            language=language,
            vad_filter=False,
            condition_on_previous_text=False,
            no_speech_threshold=no_speech,
        )
        if log_prob is not None:
            kwargs["log_prob_threshold"] = log_prob
        raw, info = model.transcribe(audio_path, **kwargs)
        segs = []
        for s in raw:
            text = (s.text or "").strip()
            if text:
                segs.append({
                    "start": round(float(s.start), 2),
                    "end":   round(float(s.end),   2),
                    "text":  text,
                })
        return segs, info

    # Pass 1
    segments, info = _run(beam=5, no_speech=0.1)

    # Pass 2 — lenient fallback for difficult languages / quiet audio
    if not segments:
        segments, info = _run(beam=1, no_speech=0.05, log_prob=-2.0)

    transcript = " ".join(s["text"] for s in segments).strip()
    det_lang = getattr(info, "language", None) or language_hint or "en"
    prob = float(getattr(info, "language_probability", 0.0) or 0.0) * 100
    duration = float(getattr(info, "duration", 0.0) or 0.0)
    return transcript, segments, det_lang, round(prob, 1), duration


# ══════════════════════════════════════════════════════════════════
#  DIARIZATION  (pause-based heuristic)
# ══════════════════════════════════════════════════════════════════
def diarize_segments(segments: List[Dict], num_speakers: int) -> List[Dict]:
    """Assign speaker labels by rotating on pauses >= 1.5s."""
    if not segments:
        return []
    out, speaker, last_end = [], 1, segments[0]["start"]
    for seg in segments:
        if num_speakers > 1 and (seg["start"] - last_end) >= 1.5:
            speaker = (speaker % num_speakers) + 1
        out.append({**seg, "speaker": f"Speaker {speaker}"})
        last_end = seg["end"]
    return out


# ══════════════════════════════════════════════════════════════════
#  GEMINI AI INTELLIGENCE  (optional — activates when user adds key)
# ══════════════════════════════════════════════════════════════════
def get_gemini_client(api_key: str):
    """Create a Gemini client from a user-supplied API key. None on failure."""
    if not api_key or not GEMINI_SDK_AVAILABLE:
        return None
    try:
        return google_genai.Client(api_key=api_key.strip())
    except Exception:
        return None


def _gemini_json(client, prompt: str, max_tokens: int = 2048) -> Optional[dict]:
    """
    Call Gemini and parse a JSON object from the response.
    Returns a dict on success, or None on any failure (caller falls back).
    """
    if client is None:
        return None
    try:
        from google.genai import types as genai_types
        config = genai_types.GenerateContentConfig(
            temperature=0.2,
            max_output_tokens=max_tokens,
            response_mime_type="application/json",
        )
        resp = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=config,
        )
        raw = (resp.text or "").strip()
        if not raw:
            return None
        # Strip code fences if present
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return _json.loads(raw)
    except Exception:
        # Retry once without the JSON mime type (some accounts differ)
        try:
            resp = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt + "\n\nReturn ONLY valid JSON, no markdown.",
            )
            raw = (resp.text or "").strip()
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)
            return _json.loads(raw)
        except Exception:
            return None


def analyze_meeting_with_gemini(client, transcript: str) -> Optional[dict]:
    """
    Use Gemini for full meeting analysis in a single call:
    summary, action items, decisions, open questions, sentiment,
    effectiveness score, and risks/blockers.

    Returns a structured dict or None if Gemini is unavailable/failed.
    """
    if client is None or not transcript.strip():
        return None

    # Cap transcript length to stay within free-tier token budget
    text = transcript[:18000]

    prompt = f"""You are an expert meeting analyst. Analyze the meeting transcript
below and return a single JSON object with EXACTLY these keys:

{{
  "summary": "A clear 3-5 sentence executive summary of what the meeting covered and its outcomes.",
  "action_items": [
    {{"text": "the task", "owner": "person responsible or 'Unspecified'", "deadline": "deadline or 'Not specified'"}}
  ],
  "decisions": [
    {{"text": "the decision that was made"}}
  ],
  "open_questions": [
    {{"text": "an unresolved question or pending item"}}
  ],
  "sentiment": {{
    "tone": "Positive | Neutral | Tense | Mixed",
    "explanation": "one sentence explaining the tone"
  }},
  "effectiveness": {{
    "score": <integer 0-100 measuring how productive and outcome-driven this meeting was>,
    "rationale": "one sentence explaining the score"
  }},
  "risks": [
    {{"text": "a risk, blocker, or concern raised (explicitly or implicitly)", "severity": "High | Medium | Low"}}
  ]
}}

Rules:
- Infer owners and deadlines from context even if not stated explicitly; use the fallbacks otherwise.
- If a section has nothing, return an empty array (or sensible defaults for sentiment/effectiveness).
- Base everything ONLY on the transcript. Do not invent facts.
- Return ONLY the JSON object, nothing else.

TRANSCRIPT:
\"\"\"{text}\"\"\"
"""

    data = _gemini_json(client, prompt, max_tokens=3072)
    if not data or "summary" not in data:
        return None
    return data


def gemini_translate(client, text: str, target_lang_name: str) -> Optional[str]:
    """Optional Gemini-based translation (more fluent than free Google Translate)."""
    if client is None or not text.strip():
        return None
    try:
        resp = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=(
                f"Translate the following text into {target_lang_name}. "
                f"Return ONLY the translation, no notes, no quotes.\n\n{text}"
            ),
        )
        out = (resp.text or "").strip()
        return out or None
    except Exception:
        return None


# ══════════════════════════════════════════════════════════════════
#  MEETING INTELLIGENCE — LIGHTWEIGHT EXTRACTORS (fallback when no Gemini key)
# ══════════════════════════════════════════════════════════════════
def _split_sentences(text: str) -> List[str]:
    return [s.strip() for s in re.split(r"(?<=[.!?])\s+", (text or "").strip())
            if len(s.strip()) > 3]


def generate_summary(transcript: str, max_sentences: int = 6) -> str:
    """Lightweight extractive summary based on meeting-relevance keywords."""
    sents = _split_sentences(transcript)
    if not sents:
        return "No transcript was generated."
    if len(sents) <= max_sentences:
        return " ".join(sents)

    keywords = [
        "decided", "agreed", "approved", "confirmed", "action", "next",
        "follow", "deadline", "task", "need", "will", "should", "review",
        "prepare", "send", "project", "issue", "risk", "plan",
        "deliverable", "timeline", "blocker", "approval",
    ]
    scored = []
    for i, s in enumerate(sents):
        lo = s.lower()
        score = sum(2 for k in keywords if k in lo) + min(len(s.split()) / 18, 2)
        if i == 0: score += 1.8
        if i == len(sents) - 1: score += 0.8
        scored.append((score, i, s))
    chosen = sorted(sorted(scored, reverse=True)[:max_sentences], key=lambda x: x[1])
    return " ".join(x[2] for x in chosen)


def _detect_deadline(text: str) -> str:
    """Extract a deadline phrase from text. Returns 'Not specified' if none found."""
    patterns = [
        r"\bby\s+(tomorrow|today|monday|tuesday|wednesday|thursday|friday|saturday|sunday)\b",
        r"\bby\s+([A-Z][a-z]+\s+\d{1,2})\b",
        r"\b(next\s+week|this\s+week|end\s+of\s+week|EOW|EOD|EOM)\b",
        r"\b\d{1,2}/\d{1,2}(?:/\d{2,4})?\b",
        r"\b(due\s+(?:on|by)?\s*[^,.!?]{3,30})",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return m.group(0).strip()
    return "Not specified"


def extract_action_items(segments: List[Dict]) -> List[Dict]:
    """Find segments that look like action items / tasks / assignments."""
    keywords = [
        "will", "should", "need to", "needs to", "must", "please",
        "send", "schedule", "review", "prepare", "follow up",
        "complete", "finish", "update", "share", "create", "submit",
        "check", "confirm", "make sure", "assign", "work on", "deliver",
    ]
    out, seen = [], set()
    for seg in segments:
        text = seg.get("text", "").strip()
        lo = text.lower()
        if len(text.split()) < 4:
            continue
        if not any(k in lo for k in keywords):
            continue
        key = re.sub(r"\W+", "", lo)[:80]
        if key in seen:
            continue
        seen.add(key)
        out.append({
            "speaker":  seg.get("speaker", "Unknown"),
            "text":     text,
            "time":     f"{seg.get('start', 0)}s",
            "owner":    seg.get("speaker", "Unknown"),
            "deadline": _detect_deadline(text),
        })
        if len(out) >= 12:
            break
    return out


def extract_decisions(segments: List[Dict]) -> List[Dict]:
    """Find segments that look like committed decisions."""
    keywords = [
        "decided", "agreed", "approved", "confirmed", "finalized",
        "accepted", "we will", "moving forward", "going with",
        "resolved", "settled", "chosen", "conclusion",
    ]
    out, seen = [], set()
    for seg in segments:
        text = seg.get("text", "").strip()
        lo = text.lower()
        if len(text.split()) < 4:
            continue
        if not any(k in lo for k in keywords):
            continue
        key = re.sub(r"\W+", "", lo)[:80]
        if key in seen:
            continue
        seen.add(key)
        out.append({
            "speaker": seg.get("speaker", "Unknown"),
            "text":    text,
            "time":    f"{seg.get('start', 0)}s",
        })
        if len(out) >= 10:
            break
    return out


def extract_open_questions(segments: List[Dict]) -> List[Dict]:
    """Find unresolved questions / pending items."""
    question_starters = (
        "what", "why", "how", "when", "where", "who",
        "can", "could", "should", "do", "does", "did", "is", "are",
    )
    open_phrases = [
        "we need to clarify", "not sure", "pending", "to be discussed",
        "need confirmation", "follow up on this", "tbd", "still unclear",
    ]
    out, seen = [], set()
    for seg in segments:
        text = seg.get("text", "").strip()
        lo = text.lower()
        first = lo.split()[0] if lo.split() else ""
        is_question = "?" in text or first in question_starters
        has_open = any(p in lo for p in open_phrases)
        if not (is_question or has_open):
            continue
        if len(text.split()) < 4:
            continue
        key = re.sub(r"\W+", "", lo)[:80]
        if key in seen:
            continue
        seen.add(key)
        out.append({
            "speaker": seg.get("speaker", "Unknown"),
            "text":    text,
            "time":    f"{seg.get('start', 0)}s",
        })
        if len(out) >= 10:
            break
    return out


# ══════════════════════════════════════════════════════════════════
#  TRANSLATION  (all 52 × 52 pairs via Google Translate)
# ══════════════════════════════════════════════════════════════════
def translate_text_safe(text: str, target_lang: str, source_lang: str = "auto") -> str:
    """
    Translate text via deep-translator's GoogleTranslator.
    Returns the original text unchanged if the library is missing or call fails.
    """
    if not text:
        return ""
    if not TRANSLATE_AVAILABLE:
        return text
    if source_lang != "auto" and source_lang == target_lang:
        return text

    tgt = GT_CODE_MAP.get(target_lang, target_lang)
    src = "auto" if source_lang == "auto" else GT_CODE_MAP.get(source_lang, source_lang)

    try:
        # Chunk long text at sentence boundaries
        chunks, buf = [], text
        while buf:
            chunk = buf[:MAX_TRANSLATE_CHARS]
            if len(buf) > MAX_TRANSLATE_CHARS:
                cut = max(chunk.rfind(". "), chunk.rfind("? "), chunk.rfind("! "))
                if cut > 300:
                    chunk = chunk[:cut + 1]
            chunks.append(chunk.strip())
            buf = buf[len(chunk):].strip()
        parts = []
        for c in chunks:
            if c:
                parts.append(GoogleTranslator(source=src, target=tgt).translate(c))
        return " ".join(p for p in parts if p)
    except Exception:
        return text


def translate_list_items(items: List[Dict], target_lang: str, source_lang: str) -> List[Dict]:
    """Translate the 'text' field of each item, storing result in 'text_tr'."""
    for item in items:
        item["text_tr"] = translate_text_safe(
            item.get("text", ""), target_lang, source_lang
        )
    return items


# ══════════════════════════════════════════════════════════════════
#  DOCX EXPORT
# ══════════════════════════════════════════════════════════════════
def _set_cell_bg(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_banner(doc: Document, title: str, color: str = "1F497D") -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_bg(cell, color)
    run = cell.paragraphs[0].add_run("  " + clean_docx(title))
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph()


def _add_metadata_table(doc: Document, meta: Dict[str, str]) -> None:
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(meta.items()):
        r0 = table.rows[i].cells[0].paragraphs[0].add_run(k)
        r0.bold = True
        table.rows[i].cells[1].paragraphs[0].add_run(clean_docx(v))


def _add_bullets(
    doc: Document, items: List[Dict], translated_label: str,
    color_hex: str, empty_text: str,
) -> None:
    if not items:
        doc.add_paragraph(empty_text)
        return
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(f"[{clean_docx(item.get('speaker', 'Unknown'))}]  ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(color_hex)
        p.add_run(clean_docx(item.get("text", "")))
        p.add_run(f"  ({clean_docx(item.get('time', ''))})").italic = True

        if item.get("deadline") and item["deadline"] != "Not specified":
            d = doc.add_paragraph()
            d.paragraph_format.left_indent = Inches(0.4)
            d.add_run("Deadline: ").bold = True
            d.add_run(clean_docx(item.get("deadline", "")))

        if item.get("text_tr") and item["text_tr"] != item.get("text", ""):
            tp = doc.add_paragraph()
            tp.paragraph_format.left_indent = Inches(0.4)
            tp.add_run(f"🌐 {translated_label}: ").font.size = Pt(9)
            tr = tp.add_run(clean_docx(item.get("text_tr", "")))
            tr.font.size = Pt(9)
            tr.italic = True


def export_docx(
    summary: str, summary_tr: str,
    actions: List[Dict], decisions: List[Dict], questions: List[Dict],
    diarized: List[Dict],
    src_lang: str, tgt_lang: str,
    num_speakers: int, word_count: int, segment_count: int,
    ai_sentiment: Optional[dict] = None,
    ai_effectiveness: Optional[dict] = None,
    ai_risks: Optional[list] = None,
) -> str:
    """Generate the professional DOCX report and return its file path."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.85)
    section.top_margin = section.bottom_margin = Inches(0.75)

    src_name = LANG_NAMES.get(src_lang, src_lang)
    tgt_name = LANG_NAMES.get(tgt_lang, tgt_lang)

    # ── Title ───────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MINUTES OF MEETING")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 73, 125)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{APP_NAME} — Multilingual Meeting Intelligence").italic = True
    doc.add_paragraph()

    # ── Metadata ────────────────────────────────────
    _add_banner(doc, "Meeting Information", "374151")
    _add_metadata_table(doc, {
        "Generated":        datetime.now().strftime("%B %d, %Y  ·  %I:%M %p"),
        "Source language":  src_name,
        "Target language":  tgt_name,
        "Speakers":         str(num_speakers),
        "Word count":       str(word_count),
        "Segments":         str(segment_count),
    })
    doc.add_paragraph()

    # ── AI Insights (Gemini only) ───────────────────
    if ai_sentiment or ai_effectiveness:
        _add_banner(doc, "AI Insights (Gemini)", "5b21b6")
        if ai_sentiment:
            p = doc.add_paragraph()
            p.add_run("Meeting tone: ").bold = True
            p.add_run(clean_docx(ai_sentiment.get("tone", "—")))
            p.add_run("  —  " + clean_docx(ai_sentiment.get("explanation", ""))).italic = True
        if ai_effectiveness:
            p = doc.add_paragraph()
            p.add_run("Effectiveness score: ").bold = True
            p.add_run(f"{ai_effectiveness.get('score', '—')}/100")
            p.add_run("  —  " + clean_docx(ai_effectiveness.get("rationale", ""))).italic = True
        if ai_risks:
            p = doc.add_paragraph()
            p.add_run("Risks & blockers:").bold = True
            for r in ai_risks:
                rp = doc.add_paragraph(style="List Bullet")
                sev = rp.add_run(f"[{clean_docx(r.get('severity', 'Medium'))}] ")
                sev.bold = True
                rp.add_run(clean_docx(r.get("text", "")))
        doc.add_paragraph()

    # ── Executive Summary ───────────────────────────
    _add_banner(doc, "1.  Executive Summary", "1F497D")
    doc.add_paragraph(clean_docx(summary or "No summary available."))
    doc.add_paragraph()

    # ── Translated Summary ──────────────────────────
    _add_banner(doc, f"2.  Translated Summary ({tgt_name})", "1F497D")
    if summary_tr and summary_tr != summary:
        p = doc.add_paragraph(clean_docx(summary_tr))
        if p.runs:
            p.runs[0].italic = True
    else:
        doc.add_paragraph("Translation unchanged or unavailable for this language pair.")
    doc.add_paragraph()

    # ── Decisions ───────────────────────────────────
    _add_banner(doc, "3.  Key Decisions", "1a5c2e")
    _add_bullets(doc, decisions, tgt_name, "1a5c2e", "No decisions were detected.")
    doc.add_paragraph()

    # ── Action Items ────────────────────────────────
    _add_banner(doc, "4.  Action Items", "8B0000")
    _add_bullets(doc, actions, tgt_name, "8B0000", "No action items were detected.")
    doc.add_paragraph()

    # ── Open Questions ──────────────────────────────
    _add_banner(doc, "5.  Open Questions", "5c3d00")
    _add_bullets(doc, questions, tgt_name, "5c3d00", "No open questions were detected.")
    doc.add_paragraph()

    # ── Full Transcript ─────────────────────────────
    _add_banner(doc, "6.  Full Transcript", "3a3a3a")
    for seg in diarized:
        p = doc.add_paragraph()
        r = p.add_run(f"[{seg.get('start', 0)}s]  {seg.get('speaker', 'Speaker')}: ")
        r.bold = True
        r.font.size = Pt(9)
        p.add_run(clean_docx(seg.get("text", ""))).font.size = Pt(9)
        if seg.get("text_tr"):
            tp = doc.add_paragraph()
            tp.paragraph_format.left_indent = Inches(0.4)
            tp.add_run(f"🌐 {tgt_name}: ").font.size = Pt(8)
            tr = tp.add_run(clean_docx(seg["text_tr"]))
            tr.font.size = Pt(8)
            tr.italic = True
    doc.add_paragraph()

    # ── Notes / Disclaimer ──────────────────────────
    _add_banner(doc, "7.  Notes & Disclaimer", "6b7280")
    disclaimer = doc.add_paragraph(
        "Speaker labels are heuristic, generated from pause-based segmentation. "
        "Translation uses a best-effort free service. For enterprise-grade accuracy, "
        "upgrade to managed APIs (OpenAI, Google Cloud Speech, AssemblyAI, Deepgram)."
    )
    if disclaimer.runs:
        disclaimer.runs[0].italic = True

    # ── Save ────────────────────────────────────────
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(TMP_DIR, f"MoM_Live_AI_Report_{ts}.docx")
    doc.save(path)
    return path


# ══════════════════════════════════════════════════════════════════
#  UI — HERO
# ══════════════════════════════════════════════════════════════════
st.markdown(f"""
<div class="hero">
  <span class="eyebrow">{APP_NAME}</span>
  <h1 class="hero-title">Minutes that <em>write themselves.</em></h1>
  <p class="hero-sub">Upload a recording. Get a polished, bilingual report with
  decisions, action items, and AI insights — in seconds.</p>
</div>
""", unsafe_allow_html=True)

c1, c2, c3, c4 = st.columns(4)
for col, icon, val, lbl in [
    (c1, "🌐", "52×52", "Languages"),
    (c2, "🎙️", "Whisper", "Transcription"),
    (c3, "🤖", "Gemini", "AI Analysis"),
    (c4, "📄", "DOCX", "Export"),
]:
    col.markdown(
        f'<div class="stat-card"><span class="stat-icon">{icon}</span>'
        f'<div class="stat-value">{val}</div>'
        f'<div class="stat-label">{lbl}</div></div>',
        unsafe_allow_html=True,
    )

st.markdown("<br>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
#  SIDEBAR
# ══════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("### Settings")
    st.markdown("---")

    # ── Gemini AI (optional, user-supplied key) ──────────────
    st.markdown("### 🤖 AI Engine")
    gemini_key = st.text_input(
        "Gemini API key (optional)",
        type="password",
        help="Paste your free Google AI Studio key to unlock AI-powered analysis. "
             "Get one at aistudio.google.com/apikey. Leave blank to use the "
             "built-in rule-based engine.",
        placeholder="AIza...",
    )
    use_gemini = bool(gemini_key.strip()) and GEMINI_SDK_AVAILABLE
    if gemini_key.strip() and not GEMINI_SDK_AVAILABLE:
        st.error("google-genai not installed — add it to requirements.txt.")
    elif use_gemini:
        st.success("✅ Gemini AI mode active")
        st.caption("Smart analysis · sentiment · effectiveness · risk detection")
    else:
        st.caption("Rule-based mode — add a key for AI-powered analysis.")
    st.markdown("---")

    target_lang = st.selectbox(
        "🌐 Target translation language",
        options=list(LANG_NAMES.keys()),
        format_func=lambda c: f"{LANG_NAMES[c]}  ({c})",
        index=list(LANG_NAMES.keys()).index("en"),
    )

    language_hint = st.selectbox(
        "🎙️ Audio language hint",
        options=["auto"] + list(LANG_NAMES.keys()),
        format_func=lambda c: "Auto-detect" if c == "auto" else f"{LANG_NAMES[c]}  ({c})",
        index=0,
        help="Leave on Auto-detect for most files; set manually if detection is wrong.",
    )

    mode = st.radio(
        "⚡ Model mode",
        options=["Fast", "Better"],
        index=0,
        horizontal=True,
        help="Fast = tiny (recommended for Streamlit Cloud).  Better = base (more accurate)."
    )
    model_size = "tiny" if mode == "Fast" else "base"

    if language_hint in NEEDS_BETTER_MODE:
        st.warning(
            f"⚠️ **{LANG_NAMES.get(language_hint, language_hint)}** transcribes better with "
            "the **Better** mode. For best accuracy, consider switching."
        )

    num_speakers = st.slider("👥 Number of speakers", 1, 6, 3,
        help="Heuristic — labels rotate on pauses ≥ 1.5s.")

    st.markdown("---")
    st.markdown("**Translation options**")
    translate_summary    = st.checkbox("Translate summary",       value=True)
    translate_actions    = st.checkbox("Translate action items",  value=True)
    translate_decisions  = st.checkbox("Translate decisions",     value=True)
    translate_questions  = st.checkbox("Translate open questions", value=True)
    translate_transcript = st.checkbox("Translate first 10 transcript lines", value=False,
        help="Slower but useful for review.")

    st.markdown("---")
    st.markdown("**Supported formats**")
    st.markdown("`mp3`   `wav`   `m4a`   `mp4`")
    st.caption(f"Max file size: {MAX_FILE_MB} MB")

    st.markdown("---")
    if TRANSLATE_AVAILABLE:
        st.success("✅ Translation engine ready")
    else:
        st.error("❌ deep-translator not installed")
    st.caption("52 × 52 = 2,704 language pairs supported")

    st.markdown("---")
    st.caption(
        "Portfolio version — lightweight local processing. "
        "Upgrade to OpenAI / Google Cloud / AssemblyAI / Deepgram for enterprise use."
    )


# ══════════════════════════════════════════════════════════════════
#  UI — UPLOAD
# ══════════════════════════════════════════════════════════════════
st.markdown(
    '<div class="sec-hdr">Upload your meeting</div>',
    unsafe_allow_html=True,
)

uploaded = st.file_uploader(
    "Drag and drop your audio or video file",
    type=SUPPORTED_EXTS,
    label_visibility="collapsed",
)

if uploaded is None:
    st.markdown("""
    <div class="empty-state">
      <span class="empty-icon">🎙️</span>
      <div class="empty-title">Drop your audio here</div>
      <div class="empty-text">Supports mp3 · wav · m4a · mp4 — any of 52 languages</div>
      <br>
      <span class="pill pill-v">Auto language detection</span>
      <span class="pill pill-p">52 languages</span>
      <span class="pill pill-g">DOCX export</span>
    </div>
    """, unsafe_allow_html=True)
    st.stop()

# Validate size
file_mb = uploaded.size / (1024 * 1024)
if file_mb > MAX_FILE_MB:
    st.error(f"❌ This file is {file_mb:.1f} MB — please upload one under {MAX_FILE_MB} MB.")
    st.stop()

# Save and display
upload_path = save_uploaded_file(uploaded)
st.audio(upload_path)

ca, cb, cc = st.columns(3)
ca.markdown(f'<span class="pill pill-v">📄 {esc(uploaded.name)}</span>',           unsafe_allow_html=True)
cb.markdown(f'<span class="pill pill-p">💾 {file_mb:.2f} MB</span>',               unsafe_allow_html=True)
cc.markdown(f'<span class="pill pill-g">🌐 → {LANG_NAMES[target_lang]}</span>',    unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
#  RENDER RESULTS
# ══════════════════════════════════════════════════════════════════
def render_results(
    transcript, segments, diarized, summary, summary_tr,
    actions, decisions, questions, det_lang, det_prob, duration,
    target_lang, docx_path,
    ai_sentiment=None, ai_effectiveness=None, ai_risks=None,
):
    det_name = LANG_NAMES.get(det_lang, det_lang)
    tgt_name = LANG_NAMES.get(target_lang, target_lang)

    # ── Results overview ────────────────────────────
    st.markdown(
        '<div class="sec-hdr">Results overview</div>',
        unsafe_allow_html=True,
    )
    r1, r2, r3, r4, r5 = st.columns(5)
    for col, val, lbl in [
        (r1, len(transcript.split()), "Words"),
        (r2, len(diarized),            "Segments"),
        (r3, len(actions),             "Action Items"),
        (r4, len(decisions),           "Decisions"),
        (r5, len(questions),           "Questions"),
    ]:
        col.markdown(
            f'<div class="result-box"><div class="result-val">{val}</div>'
            f'<div class="result-lbl">{lbl}</div></div>',
            unsafe_allow_html=True,
        )

    st.markdown("<br>", unsafe_allow_html=True)

    # ── AI Insights (Gemini only) ───────────────────
    if ai_sentiment or ai_effectiveness:
        st.markdown(
            '<div class="sec-hdr">AI insights '
            '<span style="font-size:.9rem;color:#6e6e73;font-weight:400;">'
            '· powered by Gemini</span></div>',
            unsafe_allow_html=True,
        )
        ic1, ic2 = st.columns(2)

        tone = (ai_sentiment or {}).get("tone", "—")
        tone_expl = (ai_sentiment or {}).get("explanation", "")
        tone_emoji = {"Positive": "😊", "Neutral": "😐", "Tense": "😬", "Mixed": "🤔"}.get(tone, "💬")
        ic1.markdown(
            f'<div class="result-box" style="text-align:left;padding:1.4rem">'
            f'<div style="font-size:1.6rem">{tone_emoji}</div>'
            f'<div class="result-val" style="font-size:1.8rem">{esc(tone)}</div>'
            f'<div class="result-lbl">Meeting tone</div>'
            f'<p style="margin-top:.6rem;font-size:.85rem;color:#6e6e73">{esc(tone_expl)}</p>'
            f'</div>',
            unsafe_allow_html=True,
        )

        score = (ai_effectiveness or {}).get("score", "—")
        rationale = (ai_effectiveness or {}).get("rationale", "")
        ic2.markdown(
            f'<div class="result-box" style="text-align:left;padding:1.4rem">'
            f'<div style="font-size:1.6rem">📊</div>'
            f'<div class="result-val" style="font-size:1.8rem">{esc(str(score))}<span style="font-size:1rem">/100</span></div>'
            f'<div class="result-lbl">Effectiveness score</div>'
            f'<p style="margin-top:.6rem;font-size:.85rem;color:#6e6e73">{esc(rationale)}</p>'
            f'</div>',
            unsafe_allow_html=True,
        )

        # Risks & blockers
        if ai_risks:
            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown("**Risks & blockers detected**")
            sev_style = {
                "High":   "background:#fdecea;color:#b3261e;",
                "Medium": "background:#fff4e5;color:#9a5b00;",
                "Low":    "background:#e9f6ec;color:#1c7a37;",
            }
            for r in ai_risks:
                sev = r.get("severity", "Medium")
                chip = sev_style.get(sev, sev_style["Medium"])
                st.markdown(
                    f'<div class="glass-card" style="padding:.9rem 1.1rem;margin-bottom:.5rem">'
                    f'<span class="pill" style="{chip}">{esc(sev)}</span> '
                    f'{esc(r.get("text", ""))}</div>',
                    unsafe_allow_html=True,
                )

        st.markdown("<br>", unsafe_allow_html=True)

    st.markdown(
        '<div class="sec-hdr">Executive summary</div>',
        unsafe_allow_html=True,
    )
    t1, t2 = st.tabs([f"🗣️ Original  ({det_name})", f"🌐 Translated  ({tgt_name})"])
    with t1:
        st.markdown(f'<div class="glass-card">{esc(summary)}</div>', unsafe_allow_html=True)
    with t2:
        if summary_tr == summary and det_lang != target_lang:
            st.warning("Translation may be unavailable or unchanged for this language pair.")
        st.markdown(f'<div class="glass-card">{esc(summary_tr or summary)}</div>',
                    unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Decisions ───────────────────────────────────
    st.markdown(
        '<div class="sec-hdr">Key decisions</div>',
        unsafe_allow_html=True,
    )
    if decisions:
        for i, d in enumerate(decisions, 1):
            with st.expander(f"Decision {i}  —  {d.get('speaker','')}  ({d.get('time','')})"):
                st.markdown(f"**Original:** {d.get('text','')}")
                if d.get("text_tr") and d["text_tr"] != d.get("text"):
                    st.caption(f"🌐 {tgt_name}: {d['text_tr']}")
    else:
        st.info("No explicit decisions detected.")

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Actions ─────────────────────────────────────
    st.markdown(
        '<div class="sec-hdr">Action items</div>',
        unsafe_allow_html=True,
    )
    if actions:
        for i, a in enumerate(actions, 1):
            with st.expander(f"Action {i}  —  {a.get('speaker','')}  ({a.get('time','')})"):
                st.markdown(f"**Task:** {a.get('text','')}")
                st.markdown(f"**Owner:** {a.get('owner','Unknown')}")
                st.markdown(f"**Deadline:** {a.get('deadline','Not specified')}")
                if a.get("text_tr") and a["text_tr"] != a.get("text"):
                    st.caption(f"🌐 {tgt_name}: {a['text_tr']}")
    else:
        st.info("No action items detected.")

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Open questions ──────────────────────────────
    st.markdown(
        '<div class="sec-hdr">Open questions</div>',
        unsafe_allow_html=True,
    )
    if questions:
        for i, q in enumerate(questions, 1):
            with st.expander(f"Question {i}  —  {q.get('speaker','')}  ({q.get('time','')})"):
                st.markdown(f"**Original:** {q.get('text','')}")
                if q.get("text_tr") and q["text_tr"] != q.get("text"):
                    st.caption(f"🌐 {tgt_name}: {q['text_tr']}")
    else:
        st.info("No open questions detected.")

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Transcript ──────────────────────────────────
    st.markdown(
        '<div class="sec-hdr">Full transcript</div>',
        unsafe_allow_html=True,
    )
    with st.expander(f"View transcript  ({len(diarized)} segments)"):
        for seg in diarized:
            st.markdown(
                f"**[{seg.get('start', 0)}s]  {seg.get('speaker','')}:** {seg.get('text','')}"
            )
            if seg.get("text_tr"):
                st.caption(f"🌐 {seg['text_tr']}")

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Download ────────────────────────────────────
    st.markdown(
        '<div class="sec-hdr">Download report</div>',
        unsafe_allow_html=True,
    )
    with open(docx_path, "rb") as f:
        st.download_button(
            label=f"⤓  Download Meeting Minutes  ({det_name} → {tgt_name})",
            data=f,
            file_name=os.path.basename(docx_path),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


# ══════════════════════════════════════════════════════════════════
#  GENERATE PIPELINE
# ══════════════════════════════════════════════════════════════════
if st.button("✨  Generate Meeting Minutes"):
    note = st.empty()
    note.info("Processing started — first run downloads the Whisper model, which can take a minute.")
    progress = st.progress(0)
    status = st.empty()
    t0 = time.time()

    try:
        # 1. Convert audio
        status.markdown("🎧 **Converting audio...**")
        wav_path, warn = convert_to_wav(upload_path)
        if warn:
            st.warning(f"Audio conversion note: {warn}")
        progress.progress(10)

        # 2. Transcribe
        status.markdown(f"🎙️ **Transcribing audio  ·  {mode} mode...**")
        transcript, segments, det_lang, det_prob, duration = transcribe_audio(
            wav_path, model_size, language_hint
        )
        progress.progress(38)

        if not transcript or not segments:
            status.empty(); progress.empty(); note.empty()
            st.error(
                "❌ **No speech detected.**\n\n"
                f"**Language hint:** `{language_hint}`  ·  **Mode:** `{mode}` (model: {model_size})\n\n"
                "**Likely causes:**\n"
                "- The file has no audio track\n"
                "- The audio is silent or all background music\n"
                "- The language needs the **Better** mode (Telugu, Hindi, Tamil, Arabic, etc.)\n\n"
                "Try switching to **Better** mode in the sidebar and re-uploading."
            )
            st.stop()

        # 3. Diarize
        status.markdown("👥 **Segmenting speakers...**")
        diarized = diarize_segments(segments, num_speakers)
        progress.progress(48)

        # 4-7. Meeting intelligence — Gemini if key provided, else rule-based
        gemini_data = None
        gemini_client = get_gemini_client(gemini_key) if use_gemini else None

        if gemini_client is not None:
            status.markdown("🤖 **Analyzing with Gemini AI...**")
            gemini_data = analyze_meeting_with_gemini(gemini_client, transcript)
            progress.progress(72)

        if gemini_data:
            # ── AI-powered results ──
            summary = gemini_data.get("summary", "") or generate_summary(transcript)

            actions = []
            for a in gemini_data.get("action_items", []):
                if a.get("text"):
                    actions.append({
                        "speaker":  a.get("owner", "Unspecified"),
                        "text":     a["text"],
                        "time":     "—",
                        "owner":    a.get("owner", "Unspecified"),
                        "deadline": a.get("deadline", "Not specified"),
                    })

            decisions = [{"speaker": "—", "text": d["text"], "time": "—"}
                         for d in gemini_data.get("decisions", []) if d.get("text")]

            questions = [{"speaker": "—", "text": q["text"], "time": "—"}
                         for q in gemini_data.get("open_questions", []) if q.get("text")]

            ai_sentiment     = gemini_data.get("sentiment", {})
            ai_effectiveness = gemini_data.get("effectiveness", {})
            ai_risks         = [r for r in gemini_data.get("risks", []) if r.get("text")]
        else:
            # ── Rule-based fallback ──
            if use_gemini:
                st.warning(
                    "⚠️ Gemini analysis was unavailable (check your API key or rate limits). "
                    "Falling back to the built-in engine."
                )
            status.markdown("🧠 **Generating summary...**")
            summary = generate_summary(transcript)
            progress.progress(58)

            status.markdown("📌 **Extracting action items...**")
            actions = extract_action_items(diarized)
            progress.progress(64)

            status.markdown("✅ **Extracting decisions...**")
            decisions = extract_decisions(diarized)
            progress.progress(70)

            status.markdown("❓ **Extracting open questions...**")
            questions = extract_open_questions(diarized)
            progress.progress(74)

            ai_sentiment     = None
            ai_effectiveness = None
            ai_risks         = None

        # 8. Translate (52 × 52)
        src_code = det_lang if det_lang in LANG_NAMES else "auto"
        det_name = LANG_NAMES.get(det_lang, det_lang)
        tgt_name = LANG_NAMES.get(target_lang, target_lang)
        status.markdown(f"🌐 **Translating  {det_name} → {tgt_name}...**")

        summary_tr = translate_text_safe(summary, target_lang, src_code) if translate_summary else summary
        if translate_actions:
            actions = translate_list_items(actions, target_lang, src_code)
        if translate_decisions:
            decisions = translate_list_items(decisions, target_lang, src_code)
        if translate_questions:
            questions = translate_list_items(questions, target_lang, src_code)

        if translate_transcript:
            for i, seg in enumerate(diarized):
                if i < 10:
                    seg["text_tr"] = translate_text_safe(seg.get("text", ""), target_lang, src_code)
                else:
                    seg["text_tr"] = ""
        progress.progress(88)

        # 9. DOCX
        status.markdown("📄 **Generating DOCX report...**")
        docx_path = export_docx(
            summary, summary_tr,
            actions, decisions, questions, diarized,
            det_lang, target_lang,
            num_speakers, len(transcript.split()), len(diarized),
            ai_sentiment, ai_effectiveness, ai_risks,
        )
        progress.progress(100)
        status.empty(); note.empty()

        elapsed = time.time() - t0
        st.success(f"✅ Done in {elapsed:.1f} seconds")
        st.markdown(
            f'<span class="pill pill-a">🌐 {det_name}  ·  {det_prob}%  ·  {duration:.0f}s</span>',
            unsafe_allow_html=True,
        )

        render_results(
            transcript, segments, diarized,
            summary, summary_tr,
            actions, decisions, questions,
            det_lang, det_prob, duration,
            target_lang, docx_path,
            ai_sentiment, ai_effectiveness, ai_risks,
        )

    except Exception as exc:
        status.empty(); progress.empty(); note.empty()
        st.error(f"❌ Processing failed: {str(exc)[:500]}")
        st.info(
            "Try a shorter MP3/WAV file. For Telugu/Hindi/Tamil/Arabic, switch to **Better** mode."
        )
    finally:
        gc.collect()
